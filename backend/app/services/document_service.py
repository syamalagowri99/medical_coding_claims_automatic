from sqlalchemy.orm import Session
from app.models.document import Document, DocumentType, DocumentStatus
from app.schemas.document import DocumentCreate, DocumentUpdate
from app.services.llm_service import llm_service
from app.models.entity import ExtractedEntity, EntityType
from app.models.medical_code import MedicalCode, CodeSystem, CodeStatus
from app.models.embedding import Embedding
from typing import List, Optional
import PyPDF2
import pdfplumber
import docx
from io import BytesIO


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file using pdfplumber with PyPDF2 fallback."""
    try:
        # Try pdfplumber first (better text extraction and encoding handling)
        try:
            pdf_file = BytesIO(file_content)
            with pdfplumber.open(pdf_file) as pdf:
                text = ""
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text() or ""
                        if page_text:
                            # Clean up encoding issues
                            try:
                                page_text = page_text.encode('utf-8', errors='replace').decode('utf-8')
                            except:
                                pass
                            text += page_text + "\n"
                    except Exception as page_err:
                        print(f"Warning: pdfplumber failed to extract page {page_num}: {str(page_err)}")
                        continue
                
                if text.strip():
                    return text.strip()
        except Exception as pdfplumber_err:
            print(f"pdfplumber extraction failed, trying PyPDF2: {str(pdfplumber_err)}")
        
        # Fallback to PyPDF2
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        if len(pdf_reader.pages) == 0:
            raise ValueError("PDF has no pages")
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text() or ""
                # Clean up encoding issues
                if page_text:
                    try:
                        page_text = page_text.encode('utf-8', errors='replace').decode('utf-8')
                    except:
                        pass
                text += page_text + "\n"
            except Exception as page_err:
                print(f"Warning: PyPDF2 failed to extract page {page_num}: {str(page_err)}")
                continue
        
        extracted = text.strip()
        if not extracted:
            extracted = "PDF Document: Unable to extract readable text (may be image-based or have encoding issues)"
        
        return extracted
    except Exception as e:
        print(f"Error extracting PDF text: {str(e)}")
        return f"PDF Document - Text extraction failed. Error: {str(e)[:100]}"


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    doc_file = BytesIO(file_content)
    doc = docx.Document(doc_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def create_document(
    db: Session,
    document: DocumentCreate,
    file_content: bytes,
    file_type: str,
    file_size: int,
    uploaded_by: int
) -> Document:
    """Create a new document and extract text."""
    # Extract text based on file type
    if file_type == ".pdf":
        content_text = extract_text_from_pdf(file_content)
    elif file_type == ".docx":
        content_text = extract_text_from_docx(file_content)
    else:
        content_text = file_content.decode("utf-8", errors="ignore")
    
    db_document = Document(
        patient_id=document.patient_id,
        filename=document.filename,
        file_type=file_type,
        file_size=file_size,
        document_type=document.document_type,
        status=DocumentStatus.PROCESSING,
        content_text=content_text,
        uploaded_by=uploaded_by
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    
    return db_document


def process_document(db: Session, document_id: int) -> Document:
    """Process document with LLM to extract entities and suggest codes."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise ValueError("Document not found")
    
    try:
        # Check if OpenAI API key is configured
        from app.core.config import settings
        has_api_key = settings.OPENAI_API_KEY or (settings.AZURE_OPENAI_API_KEY and settings.AZURE_OPENAI_ENDPOINT)
        
        if has_api_key:
            try:
                # Extract entities using LLM
                entities_data = llm_service.extract_entities(document.content_text)
                
                # Create extracted entities
                all_entities = []
                
                # Process diagnoses
                for diag in entities_data.get("diagnoses", []):
                    entity = ExtractedEntity(
                        document_id=document.id,
                        entity_type=EntityType.DIAGNOSIS,
                        entity_text=diag["text"],
                        context=diag.get("context"),
                        confidence_score=diag.get("confidence", 0.0)
                    )
                    db.add(entity)
                    all_entities.append(entity)
                
                # Process procedures
                for proc in entities_data.get("procedures", []):
                    entity = ExtractedEntity(
                        document_id=document.id,
                        entity_type=EntityType.PROCEDURE,
                        entity_text=proc["text"],
                        context=proc.get("context"),
                        confidence_score=proc.get("confidence", 0.0)
                    )
                    db.add(entity)
                    all_entities.append(entity)
                
                # Process medications
                for med in entities_data.get("medications", []):
                    entity = ExtractedEntity(
                        document_id=document.id,
                        entity_type=EntityType.MEDICATION,
                        entity_text=med["text"],
                        context=med.get("context"),
                        confidence_score=med.get("confidence", 0.0)
                    )
                    db.add(entity)
                    all_entities.append(entity)
                
                # Process lab tests
                for lab in entities_data.get("lab_tests", []):
                    entity = ExtractedEntity(
                        document_id=document.id,
                        entity_type=EntityType.LAB_TEST,
                        entity_text=lab["text"],
                        context=lab.get("context"),
                        confidence_score=lab.get("confidence", 0.0)
                    )
                    db.add(entity)
                    all_entities.append(entity)
                
                db.flush()
                
                # Suggest codes for each entity
                for entity in all_entities:
                    if entity.entity_type in [EntityType.DIAGNOSIS, EntityType.PROCEDURE]:
                        try:
                            suggested_codes = llm_service.suggest_codes(
                                entity.entity_text,
                                entity.entity_type.value
                            )
                            
                            for code_data in suggested_codes[:3]:  # Take top 3 suggestions
                                code_system = CodeSystem.ICD_10 if entity.entity_type == EntityType.DIAGNOSIS else CodeSystem.CPT
                                medical_code = MedicalCode(
                                    entity_id=entity.id,
                                    code_system=code_system,
                                    code=code_data["code"],
                                    description=code_data.get("description"),
                                    confidence_score=code_data.get("confidence", 0.0),
                                    status=CodeStatus.SUGGESTED
                                )
                                db.add(medical_code)
                        except Exception as code_err:
                            print(f"Warning: Failed to suggest codes for entity {entity.entity_text}: {str(code_err)}")
                            # Continue processing even if code suggestion fails
                
                # Generate embedding for the document using extraction pipeline (backend-only)
                try:
                    from app.services.extraction_pipeline import extraction_pipeline
                    embedding = extraction_pipeline.generate_document_embedding(
                        db, document.id, document.content_text
                    )
                except Exception as embed_err:
                    print(f"Warning: Failed to generate embedding for document {document_id}: {str(embed_err)}")
                    # Continue processing even if embedding fails
            
            except Exception as llm_err:
                print(f"Warning: LLM processing failed for document {document_id}: {str(llm_err)}")
                # If LLM processing fails, still mark as processed with just text extraction
                pass
        
        # Update document status to processed
        document.status = DocumentStatus.PROCESSED
        db.commit()
        db.refresh(document)
        
        return document
        
    except Exception as e:
        import traceback
        print(f"Error processing document {document_id}: {str(e)}")
        traceback.print_exc()
        document.status = DocumentStatus.FAILED
        db.commit()
        db.refresh(document)
        return document


def get_document(db: Session, document_id: int) -> Optional[Document]:
    """Get document by ID."""
    return db.query(Document).filter(Document.id == document_id).first()


def get_documents_by_patient(db: Session, patient_id: int) -> List[Document]:
    """Get all documents for a patient."""
    return db.query(Document).filter(Document.patient_id == patient_id).all()


def update_document(db: Session, document_id: int, document_update: DocumentUpdate) -> Document:
    """Update document information."""
    document = get_document(db, document_id)
    if not document:
        raise ValueError("Document not found")
    
    update_data = document_update.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(document, field, value)
    
    db.commit()
    db.refresh(document)
    return document
